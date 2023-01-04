from datetime import date
import os

from docx import Document
from docx.shared import Pt, Cm

from read_txt_report import create_case_dict, get_text


def init_document(header_txt):
    """
    This function is creating the first page of the document, with header 
    and date.

    :param header_txt: The text to add at the top of the document.
    :return: The document with the first page created.
    """
    
    document = Document()
    document.add_heading(header_txt, 0)
    document.add_heading("Date: " + str(date.today()))
    # document.add_heading("Date: " + date.today().strftime("%Y%m%d"), 1)

    return document


def document_settings(document):
    """
    This graph is setting the name, size and font of the document.

    :param document: The document to do the settings for.
    """

    font = document.styles['Normal'].font
    font.name = 'Arial'
    font.size = Pt(10)


def add_figure(document, picture_text_above, picture_text_below, picture_name, case_num):
    """
    This function is creating a new page for the specified case. On
    this page text is added above the picture, a picture with the graphs
    is added and then some text below the graph is added.
    
    :param document: The document to add the figures to.
    :param picture_text_above: The text to show above the picture.
    :param picture_text_below: The text to show below the picture. 
    :param picture_name: The picture (png file) to add to this page.
    :param case_num: The case number.
    """

    document.add_page_break()

    paragraph = document.add_paragraph('Case {}'.format(case_num) + '\n' + '\n' + picture_text_above + '\n', style = 'List Bullet')
    paragraph.style = document.styles['Normal']
    
    run = paragraph.add_run()
    run.add_picture(picture_name, width=Cm(15), height=Cm(10))
    
    paragraph = document.add_paragraph(picture_text_below + '\n', style = 'List Bullet')
    paragraph.style = document.styles['Normal']


def save_document(document, file_path, file_name):
    """
    This function saves the document under the specified file_name.
    
    :param document: The document to save.
    :param file_path: The path to where the file should be saved.
    :param file_name: The name of the file, when it is saved.
    """

    document.save(os.path.join(file_path, file_name + '.docx'))  
    
    return


def create_document(header_text, num_cases, txt_file, path_png_files, project_name, save_path, document_name):
    """
    This function creates the document, runs through all the cases
    in order to add the graphs for each case and saves the document at the end.
    
    :param header_text: The text at the top of the document - the title.
    :param num_cases: The number of cases to add to the document.
    :param txt_file: The name of the text file with the text for the document.
    :param path_png_files: The path to where the png files are placed.
    :param project_name: The name of the project to create the document for.
    :param save_path: The path to where the generated document should be saved.
    :param document_name: The name of the generated document when saved.
    """

    document = init_document(header_txt=header_text)
    
    document_settings(document)
    
    case_dict = create_case_dict(txt_file)

    for case_num in range(1, num_cases + 1):

        text_above_picture, text_below_picture = get_text(case_num, case_dict)
        
        name_of_picture = os.path.join(path_png_files, project_name + '_0{}.png'.format(case_num)) if case_num < 10 else os.path.join(path_png_files, project_name + '_{}.png'.format(case_num))
        
        add_figure(
            document=document,
            picture_text_above=text_above_picture,
            picture_text_below=text_below_picture,
            picture_name=name_of_picture,
            case_num=case_num
            )

    save_document(document, save_path, document_name)



    


